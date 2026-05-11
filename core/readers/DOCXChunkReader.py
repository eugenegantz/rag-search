import docx
import typing
from docx.document import Document as DocumentObject
from core.chunking import create_chunks_with_coords
from core.readers.BaseChunkReader import BaseChunkReader
from core.types import TChunk


class DOCXChunkReader(BaseChunkReader):
    """Ридер для DOCX-файлов. Координаты чанка: [paragraph, char_idx]."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self._doc: DocumentObject | None = None


    def load(self) -> None:
        if self._doc is None:
            self._doc = docx.Document(self.filepath)


    def is_loaded(self) -> bool:
        return self._doc is not None


    def _chars_with_coords(self):
        """Генератор символов с координатами [paragraph, char_idx] без накопления в памяти."""
        self.load()

        if not self._doc:
            return

        for para_idx, para in enumerate(self._doc.paragraphs):
            for char_idx, char in enumerate(para.text):
                yield char, (para_idx, char_idx)
            # Абзацы не имеют в конце переноса строки -- их необходимо разделить явно
            yield "\n", (para_idx, len(para.text) + 0)
            yield "\n", (para_idx, len(para.text) + 1)


    def getChunk(self, from_: list[int], to: list[int]) -> TChunk:
        self.load()

        if not self._doc:
            raise Exception("doc is None")

        start_para, start_char  = from_
        end_para, end_char      = to
        texts: list[str]        = []

        for para_idx in range(start_para, end_para + 1):
            para = self._doc.paragraphs[para_idx]
            text = para.text

            if para_idx == start_para and para_idx == end_para:
                texts.append(text[start_char:end_char + 1])

            elif para_idx == start_para:
                texts.append(text[start_char:])

            elif para_idx == end_para:
                texts.append(text[:end_char + 1])

            else:
                texts.append(text)

        return {
            "text": " ".join(texts),
            "from": from_,
            "to": to,
        }


    def createChunksIterator(self) -> typing.Iterator[TChunk]:
        return create_chunks_with_coords(
            self._chars_with_coords(),
            delimiters=[" "],
        )


    def createChunks(self) -> list[TChunk]:
        return [chunk for chunk in self.createChunksIterator()]
